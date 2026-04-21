"""文档处理工具 —— PDF、DOCX 与富预览能力。

在基础文件工具之上补充了特定格式的读取器。
可选依赖包括：pymupdf（PDF）、python-docx（DOCX）、chardet（编码检测）。
"""
from __future__ import annotations

import mimetypes
from pathlib import Path

from scaffold.tools.errors import ToolError, ToolErrorCode

from fs_agent.tools.file_tools import registry, _check_sandbox


# ---------------------------------------------------------------------------
# 工具：read_pdf
# ---------------------------------------------------------------------------

@registry.tool
def read_pdf(path: str, pages: str = "all", max_pages: int = 50) -> str:
    """读取 PDF 文件中的文本内容，适用于 `.pdf` 文件。

    对 PDF 应优先使用它而不是 `read_file`，因为 `read_file` 返回的是原始字节。

    path: PDF 文件路径。
    pages: 要读取的页码。`all` 表示全部，`1-5` 表示范围，`1,3,5` 表示指定页。
    max_pages: 最多提取的页数（安全限制）。
    """
    resolved = _check_sandbox(path)
    if not resolved.is_file():
        raise ToolError(ToolErrorCode.NOT_FOUND, f"'{path}' is not a file.")
    if resolved.suffix.lower() != ".pdf":
        raise ToolError(ToolErrorCode.UNSUPPORTED_FORMAT, f"'{path}' is not a PDF file.")

    try:
        import fitz  # pymupdf 库
    except ImportError:
        raise ToolError(
            ToolErrorCode.INTERNAL,
            "PDF support requires pymupdf. Install with: pip install pymupdf",
        )

    try:
        doc = fitz.open(str(resolved))
    except Exception as e:
        raise ToolError(ToolErrorCode.INTERNAL, f"Failed to open PDF: {e}")

    total_pages = len(doc)
    page_indices = _parse_page_range(pages, total_pages)
    if len(page_indices) > max_pages:
        page_indices = page_indices[:max_pages]

    parts: list[str] = []
    parts.append(f"[PDF: {resolved.name} | {total_pages} pages | extracting {len(page_indices)} page(s)]")

    for idx in page_indices:
        page = doc[idx]
        text = page.get_text()
        if text.strip():
            parts.append(f"\n--- Page {idx + 1} ---\n{text.strip()}")
        else:
            parts.append(f"\n--- Page {idx + 1} ---\n(no extractable text)")

    doc.close()
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# 工具：read_docx
# ---------------------------------------------------------------------------

@registry.tool
def read_docx(path: str, max_paragraphs: int = 500) -> str:
    """读取 Word 文档（`.docx`）中的文本内容，适用于 `.docx` 文件。

    path: Word 文档路径。
    max_paragraphs: 最多提取的段落数。
    """
    resolved = _check_sandbox(path)
    if not resolved.is_file():
        raise ToolError(ToolErrorCode.NOT_FOUND, f"'{path}' is not a file.")
    if resolved.suffix.lower() != ".docx":
        raise ToolError(ToolErrorCode.UNSUPPORTED_FORMAT, f"'{path}' is not a .docx file.")

    try:
        from docx import Document
    except ImportError:
        raise ToolError(
            ToolErrorCode.INTERNAL,
            "DOCX support requires python-docx. Install with: pip install python-docx",
        )

    try:
        doc = Document(str(resolved))
    except Exception as e:
        raise ToolError(ToolErrorCode.INTERNAL, f"Failed to open document: {e}")

    paragraphs = []
    for i, para in enumerate(doc.paragraphs):
        if i >= max_paragraphs:
            break
        text = para.text.strip()
        if text:
            # 保留标题结构
            if para.style and para.style.name.startswith("Heading"):
                level = para.style.name.replace("Heading ", "")
                try:
                    hashes = "#" * int(level)
                except ValueError:
                    hashes = "#"
                paragraphs.append(f"{hashes} {text}")
            else:
                paragraphs.append(text)

    # 同时提取表格内容
    tables_text = []
    for t_idx, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            tables_text.append(f"\n[Table {t_idx + 1}]\n" + "\n".join(rows))

    header = f"[DOCX: {resolved.name} | {len(doc.paragraphs)} paragraphs | {len(doc.tables)} tables]"
    body = "\n\n".join(paragraphs[:max_paragraphs])
    tables = "\n".join(tables_text) if tables_text else ""

    return f"{header}\n\n{body}{tables}"


# ---------------------------------------------------------------------------
# 工具：preview_file
# ---------------------------------------------------------------------------

@registry.tool
def preview_file(path: str) -> str:
    """获取任意文件的智能预览，自动检测格式并提取关键信息。

    当你不了解文件格式，只想快速看一眼时可以使用它。
    支持：文本、PDF、DOCX、代码、CSV、JSON、YAML。

    path: 需要预览的文件路径。
    """
    resolved = _check_sandbox(path)
    if not resolved.is_file():
        raise ToolError(ToolErrorCode.NOT_FOUND, f"'{path}' is not a file.")

    size = resolved.stat().st_size
    suffix = resolved.suffix.lower()

    # 分发到对应的专用读取器
    if suffix == ".pdf":
        return read_pdf(path, pages="1-3", max_pages=3)
    elif suffix == ".docx":
        return read_docx(path, max_paragraphs=30)
    elif suffix in (".csv", ".tsv"):
        return _preview_csv(resolved, suffix)
    elif suffix in (".json", ".yaml", ".yml"):
        return _preview_structured(resolved)
    else:
        return _preview_text(resolved, size)


# ---------------------------------------------------------------------------
# 工具：summarize_file
# ---------------------------------------------------------------------------

@registry.tool
def summarize_file(path: str) -> str:
    """提取文件的关键结构信息，帮助快速理解内容。

    返回内容包括：文件类型、大小、结构信息（例如文档标题、代码函数名等）。
    它不会读取完整内容；若需全文，请使用 `read_file` 或 `read_pdf`。

    path: 文件路径。
    """
    resolved = _check_sandbox(path)
    if not resolved.is_file():
        raise ToolError(ToolErrorCode.NOT_FOUND, f"'{path}' is not a file.")

    size = resolved.stat().st_size
    suffix = resolved.suffix.lower()
    mime, _ = mimetypes.guess_type(str(resolved))

    info: list[str] = [
        f"File: {resolved.name}",
        f"Size: {_human_size(size)}",
        f"Type: {mime or 'unknown'} ({suffix or 'no extension'})",
    ]

    if suffix == ".pdf":
        try:
            import fitz
            doc = fitz.open(str(resolved))
            info.append(f"Pages: {len(doc)}")
            toc = doc.get_toc()
            if toc:
                info.append("Table of Contents:")
                for level, title, page in toc[:20]:
                    indent = "  " * (level - 1)
                    info.append(f"  {indent}{title} (p.{page})")
            doc.close()
        except ImportError:
            info.append("(PDF details require pymupdf)")
    elif suffix == ".docx":
        try:
            from docx import Document
            doc = Document(str(resolved))
            info.append(f"Paragraphs: {len(doc.paragraphs)}")
            info.append(f"Tables: {len(doc.tables)}")
            headings = [p.text for p in doc.paragraphs
                        if p.style and p.style.name.startswith("Heading")]
            if headings:
                info.append("Headings:")
                for h in headings[:15]:
                    info.append(f"  - {h}")
        except ImportError:
            info.append("(DOCX details require python-docx)")
    elif suffix in (".py", ".js", ".ts", ".java", ".go", ".rs"):
        info.extend(_summarize_code(resolved))
    elif suffix in (".csv", ".tsv"):
        info.extend(_summarize_csv(resolved, suffix))

    return "\n".join(info)


# ---------------------------------------------------------------------------
# 内部辅助函数
# ---------------------------------------------------------------------------

def _parse_page_range(spec: str, total: int) -> list[int]:
    """将页码规格解析为从 0 开始的索引列表。"""
    if spec.strip().lower() == "all":
        return list(range(total))
    indices = []
    for part in spec.split(","):
        part = part.strip()
        if "-" in part:
            a, b = part.split("-", 1)
            start = max(0, int(a) - 1)
            end = min(total, int(b))
            indices.extend(range(start, end))
        else:
            idx = int(part) - 1
            if 0 <= idx < total:
                indices.append(idx)
    return indices


def _preview_csv(path: Path, suffix: str) -> str:
    import csv
    delimiter = "\t" if suffix == ".tsv" else ","
    lines: list[str] = []
    with open(path, newline="", encoding="utf-8", errors="replace") as f:
        reader = csv.reader(f, delimiter=delimiter)
        for i, row in enumerate(reader):
            if i >= 20:
                lines.append("... (showing first 20 rows)")
                break
            lines.append(" | ".join(row))
    header = f"[CSV: {path.name} | delimiter='{delimiter}']"
    return header + "\n" + "\n".join(lines)


def _preview_structured(path: Path) -> str:
    text = path.read_text(encoding="utf-8", errors="replace")
    # 预览最多截断到 3000 个字符
    if len(text) > 3000:
        return f"[{path.suffix.upper()}: {path.name} | {len(text):,} chars]\n{text[:3000]}\n... (truncated)"
    return f"[{path.suffix.upper()}: {path.name} | {len(text):,} chars]\n{text}"


def _preview_text(path: Path, size: int) -> str:
    try:
        # 检测编码
        try:
            import chardet
            raw = path.read_bytes()[:10000]
            detected = chardet.detect(raw)
            encoding = detected.get("encoding", "utf-8") or "utf-8"
        except ImportError:
            encoding = "utf-8"

        text = path.read_text(encoding=encoding, errors="replace")
        preview = text[:3000]
        header = f"[Text: {path.name} | {_human_size(size)} | encoding={encoding}]"
        if len(text) > 3000:
            return f"{header}\n{preview}\n... (truncated, {len(text):,} total chars)"
        return f"{header}\n{preview}"
    except Exception:
        return f"[Binary file: {path.name} | {_human_size(size)}] (cannot preview)"


def _summarize_code(path: Path) -> list[str]:
    """从源代码中提取函数名和类名。"""
    info: list[str] = []
    try:
        text = path.read_text(encoding="utf-8", errors="replace")
        lines = text.splitlines()
        info.append(f"Lines: {len(lines)}")

        functions = []
        classes = []
        for line in lines:
            stripped = line.strip()
            if stripped.startswith("def "):
                name = stripped.split("(")[0].replace("def ", "")
                functions.append(name)
            elif stripped.startswith("class "):
                name = stripped.split("(")[0].split(":")[0].replace("class ", "")
                classes.append(name)

        if classes:
            info.append(f"Classes: {', '.join(classes[:10])}")
        if functions:
            info.append(f"Functions: {', '.join(functions[:15])}")
    except Exception:
        pass
    return info


def _summarize_csv(path: Path, suffix: str) -> list[str]:
    import csv
    info: list[str] = []
    delimiter = "\t" if suffix == ".tsv" else ","
    try:
        with open(path, newline="", encoding="utf-8", errors="replace") as f:
            reader = csv.reader(f, delimiter=delimiter)
            header = next(reader, None)
            row_count = sum(1 for _ in reader) + (1 if header else 0)
            if header:
                info.append(f"Columns: {', '.join(header)}")
            info.append(f"Rows: {row_count}")
    except Exception:
        pass
    return info


def _human_size(size: int) -> str:
    for unit in ("B", "KB", "MB", "GB"):
        if size < 1024:
            return f"{size:.1f} {unit}"
        size /= 1024  # type: ignore  # 这里接受浮点递减
    return f"{size:.1f} TB"
